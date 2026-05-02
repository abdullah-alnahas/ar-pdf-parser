"""Phase-7 docx emitter structural tests."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest

from arabic_pdf_transcribe.emit import emit_docx
from arabic_pdf_transcribe.regions import (
    BBox,
    ListMarker,
    Region,
    RegionRole,
    RegionSource,
    TableCell,
    TableGrid,
)


def _region(
    *,
    role: RegionRole = RegionRole.PARAGRAPH,
    text: str = "",
    page_index: int = 0,
    heading_level: int | None = None,
    list_marker: ListMarker | None = None,
    table_grid: TableGrid | None = None,
    group_id: str | None = None,
    failure_reason: str | None = None,
) -> Region:
    return Region(
        page_index=page_index,
        bbox=BBox(0.0, 0.0, 100.0, 30.0),
        text=text,
        role=role,
        source=RegionSource.NATIVE,
        heading_level=heading_level,
        list_marker=list_marker,
        table_grid=table_grid,
        group_id=group_id,
        failure_reason=failure_reason,
    )


def _open(path: Path):
    from docx import Document

    return Document(str(path))


# ---------------------------------------------------------------------------
# Per-role style assertions
# ---------------------------------------------------------------------------


def test_paragraph_uses_normal_style(tmp_path: Path) -> None:
    out = tmp_path / "out.docx"
    emit_docx([_region(text="hi")], out)
    doc = _open(out)
    paras = list(doc.paragraphs)
    assert len(paras) == 1
    assert paras[0].text == "hi"
    assert paras[0].style.name == "Normal"


def test_heading_uses_heading_style_with_level(tmp_path: Path) -> None:
    out = tmp_path / "h.docx"
    regions = [
        _region(role=RegionRole.HEADING, text="A", heading_level=1),
        _region(role=RegionRole.HEADING, text="B", heading_level=2),
        _region(role=RegionRole.HEADING, text="C", heading_level=3),
    ]
    emit_docx(regions, out)
    doc = _open(out)
    styles = [p.style.name for p in doc.paragraphs]
    assert "Heading 1" in styles
    assert "Heading 2" in styles
    assert "Heading 3" in styles


def test_heading_default_level_is_2(tmp_path: Path) -> None:
    out = tmp_path / "h2.docx"
    emit_docx(
        [_region(role=RegionRole.HEADING, text="X", heading_level=None)],
        out,
    )
    doc = _open(out)
    [p] = list(doc.paragraphs)
    assert p.style.name == "Heading 2"


def test_bullet_list_item_uses_list_bullet_style(tmp_path: Path) -> None:
    out = tmp_path / "l.docx"
    r = _region(
        role=RegionRole.LIST_ITEM,
        text="- one",
        list_marker=ListMarker(kind="bullet", raw_marker="-"),
    )
    emit_docx([r], out)
    doc = _open(out)
    [p] = list(doc.paragraphs)
    assert p.style.name == "List Bullet"
    # Marker stripped before emission.
    assert p.text == "one"


def test_ordered_list_item_uses_list_number_style(tmp_path: Path) -> None:
    out = tmp_path / "o.docx"
    r = _region(
        role=RegionRole.LIST_ITEM,
        text="3) third",
        list_marker=ListMarker(kind="ordered", ordinal=3, raw_marker="3)"),
    )
    emit_docx([r], out)
    doc = _open(out)
    [p] = list(doc.paragraphs)
    assert p.style.name == "List Number"


def test_caption_paragraph_is_italic(tmp_path: Path) -> None:
    out = tmp_path / "c.docx"
    emit_docx([_region(role=RegionRole.CAPTION, text="cap")], out)
    doc = _open(out)
    [p] = list(doc.paragraphs)
    assert p.text == "cap"
    [run] = p.runs
    assert run.italic is True


def test_failure_uses_quote_style(tmp_path: Path) -> None:
    out = tmp_path / "f.docx"
    r = _region(
        role=RegionRole.FAILURE_PLACEHOLDER,
        page_index=4,
        failure_reason="ocr_failed",
    )
    emit_docx([r], out)
    doc = _open(out)
    [p] = list(doc.paragraphs)
    assert p.style.name == "Quote"
    assert "page 5" in p.text
    assert "ocr_failed" in p.text


def test_header_footer_is_suppressed(tmp_path: Path) -> None:
    out = tmp_path / "hf.docx"
    emit_docx([_region(role=RegionRole.HEADER_FOOTER, text="page 1")], out)
    doc = _open(out)
    assert len(list(doc.paragraphs)) == 0


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------


def _grid(rows: list[list[str]]) -> TableGrid:
    return TableGrid(
        rows=tuple(
            tuple(TableCell(text=cell, confidence=None, bbox=BBox(0, 0, 10, 10)) for cell in row)
            for row in rows
        )
    )


def test_table_emits_word_table(tmp_path: Path) -> None:
    out = tmp_path / "t.docx"
    grid = _grid([["h1", "h2"], ["a", "b"]])
    emit_docx([_region(role=RegionRole.TABLE, table_grid=grid)], out)
    doc = _open(out)
    [table] = doc.tables
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    assert table.cell(0, 0).text == "h1"
    assert table.cell(1, 1).text == "b"


def test_table_short_row_padded(tmp_path: Path) -> None:
    out = tmp_path / "t2.docx"
    grid = _grid([["a", "b", "c"], ["x"]])
    emit_docx([_region(role=RegionRole.TABLE, table_grid=grid)], out)
    doc = _open(out)
    [table] = doc.tables
    assert len(table.columns) == 3
    assert table.cell(1, 1).text == ""
    assert table.cell(1, 2).text == ""


# ---------------------------------------------------------------------------
# Figures + grouping
# ---------------------------------------------------------------------------


def test_figure_without_group_emits_placeholder(tmp_path: Path) -> None:
    out = tmp_path / "fig.docx"
    emit_docx([_region(role=RegionRole.FIGURE, page_index=2)], out)
    doc = _open(out)
    [p] = list(doc.paragraphs)
    assert p.text == "Figure on page 3"


def test_orphan_grouped_caption_still_renders(tmp_path: Path) -> None:
    """Regression: grouped caption without a matching figure must
    still be written as a caption paragraph."""
    out = tmp_path / "orphan.docx"
    cap = _region(role=RegionRole.CAPTION, text="orphan", group_id="solo")
    emit_docx([cap], out)
    doc = _open(out)
    [p] = list(doc.paragraphs)
    assert p.text == "orphan"
    [run] = p.runs
    assert run.italic is True


def test_grouped_figure_caption_combined(tmp_path: Path) -> None:
    out = tmp_path / "fig2.docx"
    fig = _region(role=RegionRole.FIGURE, page_index=0, group_id="g1")
    cap = _region(role=RegionRole.CAPTION, text="A horse", group_id="g1")
    emit_docx([fig, cap], out)
    doc = _open(out)
    paras = [p.text for p in doc.paragraphs]
    # Caption suppressed, figure paragraph carries the caption text.
    assert paras == ["Figure on page 1: A horse"]


# ---------------------------------------------------------------------------
# Network purity
# ---------------------------------------------------------------------------


def test_emit_docx_performs_no_network_io(tmp_path: Path) -> None:
    """Acceptance criterion: emitter performs no network I/O.

    We patch ``socket.socket`` to raise on instantiation; if the
    emitter or any dep tries to create a socket the test fails.
    """
    out = tmp_path / "net.docx"
    regions = [
        _region(role=RegionRole.HEADING, text="t", heading_level=1),
        _region(text="prose"),
    ]

    def boom(*args: object, **kwargs: object) -> None:
        raise RuntimeError("network access forbidden during emit")

    with patch("socket.socket", new=boom):
        emit_docx(regions, out)
    assert out.exists()


# ---------------------------------------------------------------------------
# Determinism
# ---------------------------------------------------------------------------


def test_two_runs_produce_byte_identical_xml(tmp_path: Path) -> None:
    """python-docx writes timestamps inside core.xml, but document.xml
    (the body) should be byte-identical across runs for the same
    input. We check document.xml only.
    """
    import zipfile

    regions = [
        _region(role=RegionRole.HEADING, text="t", heading_level=1),
        _region(text="prose"),
    ]
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    emit_docx(regions, a)
    emit_docx(regions, b)
    with zipfile.ZipFile(a) as za, zipfile.ZipFile(b) as zb:
        assert za.read("word/document.xml") == zb.read("word/document.xml")


# ---------------------------------------------------------------------------
# String path acceptance
# ---------------------------------------------------------------------------


def test_string_path_accepted(tmp_path: Path) -> None:
    out = str(tmp_path / "s.docx")
    emit_docx([_region(text="x")], out)
    assert Path(out).exists()


def test_empty_input_writes_empty_doc(tmp_path: Path) -> None:
    out = tmp_path / "e.docx"
    emit_docx([], out)
    assert out.exists()
    doc = _open(out)
    assert len(list(doc.paragraphs)) == 0


# ---------------------------------------------------------------------------
# Pyright import smoke (no test logic, just ensures module loads)
# ---------------------------------------------------------------------------


def test_emit_docx_callable() -> None:
    assert callable(emit_docx)


@pytest.mark.parametrize("level", [1, 2, 3, 4, 5, 6, 9])
def test_heading_levels_clamp_to_word_max(tmp_path: Path, level: int) -> None:
    out = tmp_path / f"h{level}.docx"
    r = _region(role=RegionRole.HEADING, text="X", heading_level=level)
    emit_docx([r], out)
    doc = _open(out)
    [p] = list(doc.paragraphs)
    assert p.style.name == f"Heading {level}"
