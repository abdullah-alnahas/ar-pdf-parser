"""Word (.docx) emitter.

Converts an ordered, role-classified :class:`Region` stream into a
``python-docx`` document. The mapping mirrors :mod:`emit.markdown`
but uses Word's built-in styles instead of Markdown syntax.

Mapping (see plan section "Phase 7"):

* ``HEADING`` → built-in style ``Heading {level}``.
* ``PARAGRAPH`` → ``Normal``.
* ``LIST_ITEM`` → ``List Bullet`` or ``List Number``.
* ``TABLE`` → real Word table from ``region.table_grid``; one
  paragraph per cell.
* ``FIGURE`` / ``CAPTION`` → placeholder paragraph
  ``"Figure on page N"``; grouped caption text is appended.
* ``HEADER_FOOTER`` → suppressed.
* ``FAILURE_PLACEHOLDER`` → ``Quote`` style paragraph.
* ``UNKNOWN`` → ``Normal``.

When ``rtl=True`` (default — Arabic-first), every paragraph receives
``<w:bidi/>`` + right-aligned, every run gets ``<w:rtl/>``, and tables
get ``<w:bidiVisual/>`` so columns lay out right-to-left.

When ``page_breaks=True`` (default), an explicit page break separates
regions belonging to different source PDF pages so the Word output
mirrors the original pagination.

The emitter is **deterministic**: no timestamps, no environment
data, no document author / company metadata.

``python-docx`` is imported lazily — importing this module costs
nothing until :func:`emit_docx` is first called.
"""

from __future__ import annotations

from collections.abc import Iterable, Sequence
from pathlib import Path
from typing import TYPE_CHECKING

from arabic_pdf_transcribe.emit._bidi import add_rlm_if_needed
from arabic_pdf_transcribe.emit._normalise import (
    DEFAULT_FORM,
    NormalisationForm,
    normalise_text,
)
from arabic_pdf_transcribe.regions import (
    Region,
    RegionRole,
    TableGrid,
)

if TYPE_CHECKING:  # pragma: no cover — typing-only
    from docx.document import Document as _DocxDocument
    from docx.table import Table as _DocxTable
    from docx.text.paragraph import Paragraph as _DocxParagraph
    from docx.text.run import Run as _DocxRun


def emit_docx(
    regions: Iterable[Region],
    output_path: Path | str,
    *,
    normalisation: NormalisationForm = DEFAULT_FORM,
    apply_bidi: bool = True,
    rtl: bool = True,
    page_breaks: bool = True,
) -> None:
    """Render ``regions`` to a Word file at ``output_path``.

    Parameters
    ----------
    rtl:
        When ``True`` (default), every paragraph and table is marked
        right-to-left; runs receive ``<w:rtl/>``. Set ``False`` to emit
        an LTR document.
    page_breaks:
        When ``True`` (default), a page break is inserted between
        regions whose ``page_index`` differs, so each source PDF page
        lands on its own Word page.

    The output file is overwritten if it exists. The function does
    not perform any network I/O.
    """
    from docx import Document

    document = Document()
    region_list = list(regions)
    caption_group_ids = _collect_paired_caption_group_ids(region_list)
    pending_caption_text: dict[str, str] = _collect_caption_texts(
        region_list, normalisation=normalisation, paired=caption_group_ids
    )

    last_page_index: int | None = None
    pending_break = False
    for region in region_list:
        role = region.role
        if role is RegionRole.HEADER_FOOTER:
            continue
        if role is RegionRole.CAPTION and region.group_id in caption_group_ids:
            # Already rendered alongside the figure.
            continue
        if (
            page_breaks
            and last_page_index is not None
            and region.page_index != last_page_index
        ):
            pending_break = True
        last_page_index = region.page_index

        if role is RegionRole.HEADING:
            paragraph = _add_heading(
                document,
                region,
                normalisation=normalisation,
                apply_bidi=apply_bidi,
            )
        elif role is RegionRole.LIST_ITEM:
            paragraph = _add_list_item(
                document,
                region,
                normalisation=normalisation,
                apply_bidi=apply_bidi,
            )
        elif role is RegionRole.TABLE:
            if pending_break:
                # Tables don't carry pPr we control directly; emit a
                # marker paragraph that owns the page-break-before.
                _add_page_break_anchor(document, rtl=rtl)
                pending_break = False
            _add_table(document, region, normalisation=normalisation, rtl=rtl)
            continue  # tables don't go through the paragraph RTL path
        elif role is RegionRole.FIGURE:
            paragraph = _add_figure(
                document,
                region,
                pending_caption_text=pending_caption_text,
            )
        elif role is RegionRole.CAPTION:
            paragraph = _add_caption(
                document,
                region,
                normalisation=normalisation,
                apply_bidi=apply_bidi,
            )
        elif role is RegionRole.FAILURE_PLACEHOLDER:
            paragraph = _add_failure(document, region)
        else:  # PARAGRAPH / UNKNOWN
            paragraph = _add_paragraph(
                document,
                region,
                normalisation=normalisation,
                apply_bidi=apply_bidi,
            )

        if pending_break:
            _set_page_break_before(paragraph)
            pending_break = False
        if rtl:
            _set_rtl_paragraph(paragraph)

    document.save(str(output_path))


# ---------------------------------------------------------------------------
# Pre-pass collectors
# ---------------------------------------------------------------------------


def _collect_paired_caption_group_ids(regions: Sequence[Region]) -> set[str]:
    """Group ids where BOTH a caption and figure exist.

    Captions are suppressed only when a paired figure consumes their
    text; orphan grouped captions still render so the text is not
    dropped.
    """
    figure_ids = {
        region.group_id
        for region in regions
        if region.role is RegionRole.FIGURE and region.group_id is not None
    }
    caption_ids = {
        region.group_id
        for region in regions
        if region.role is RegionRole.CAPTION and region.group_id is not None
    }
    return figure_ids & caption_ids


def _collect_caption_texts(
    regions: Sequence[Region],
    *,
    normalisation: NormalisationForm,
    paired: set[str],
) -> dict[str, str]:
    """Return ``group_id → caption-text`` for paired (figure+caption) groups."""
    out: dict[str, str] = {}
    for region in regions:
        if (
            region.role is RegionRole.CAPTION
            and region.group_id is not None
            and region.group_id in paired
        ):
            out[region.group_id] = normalise_text(region.text, form=normalisation).strip()
    return out


# ---------------------------------------------------------------------------
# RTL / page-break helpers
# ---------------------------------------------------------------------------


def _set_rtl_paragraph(paragraph: _DocxParagraph) -> None:
    """Mark ``paragraph`` right-to-left (bidi + right-aligned + rtl runs)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))
    for run in paragraph.runs:
        _set_rtl_run(run)


def _set_rtl_run(run: _DocxRun) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rPr = run._r.get_or_add_rPr()
    if rPr.find(qn("w:rtl")) is None:
        rPr.append(OxmlElement("w:rtl"))


def _set_rtl_table(table: _DocxTable) -> None:
    """Mark ``table`` right-to-left layout + RTL paragraphs in every cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tblPr = table._tbl.tblPr
    if tblPr is not None and tblPr.find(qn("w:bidiVisual")) is None:
        tblPr.append(OxmlElement("w:bidiVisual"))
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _set_rtl_paragraph(paragraph)


def _set_page_break_before(paragraph: _DocxParagraph) -> None:
    paragraph.paragraph_format.page_break_before = True


def _add_page_break_anchor(document: _DocxDocument, *, rtl: bool) -> None:
    """Empty paragraph that just owns a page-break-before flag.

    Used before tables, which don't expose ``page_break_before`` themselves.
    """
    paragraph = document.add_paragraph(style="Normal")
    _set_page_break_before(paragraph)
    if rtl:
        _set_rtl_paragraph(paragraph)


# ---------------------------------------------------------------------------
# Per-role writers
# ---------------------------------------------------------------------------


def _prepare_text(text: str, *, normalisation: NormalisationForm, apply_bidi: bool) -> str:
    out = normalise_text(text, form=normalisation)
    if apply_bidi:
        out = add_rlm_if_needed(out)
    return out


def _add_heading(
    document: _DocxDocument,
    region: Region,
    *,
    normalisation: NormalisationForm,
    apply_bidi: bool,
) -> _DocxParagraph:
    level = region.heading_level if region.heading_level is not None else 2
    level = max(1, min(level, 9))
    text = _prepare_text(region.text, normalisation=normalisation, apply_bidi=apply_bidi)
    return document.add_paragraph(text, style=f"Heading {level}")


def _add_paragraph(
    document: _DocxDocument,
    region: Region,
    *,
    normalisation: NormalisationForm,
    apply_bidi: bool,
) -> _DocxParagraph:
    text = _prepare_text(region.text, normalisation=normalisation, apply_bidi=apply_bidi)
    return document.add_paragraph(text, style="Normal")


def _add_caption(
    document: _DocxDocument,
    region: Region,
    *,
    normalisation: NormalisationForm,
    apply_bidi: bool,
) -> _DocxParagraph:
    text = _prepare_text(region.text, normalisation=normalisation, apply_bidi=apply_bidi)
    paragraph = document.add_paragraph(style="Normal")
    run = paragraph.add_run(text)
    run.italic = True
    return paragraph


def _add_list_item(
    document: _DocxDocument,
    region: Region,
    *,
    normalisation: NormalisationForm,
    apply_bidi: bool,
) -> _DocxParagraph:
    marker = region.list_marker
    raw_text = region.text
    if marker is not None and marker.raw_marker:
        stripped = raw_text.lstrip()
        if stripped.startswith(marker.raw_marker):
            leading_ws = raw_text[: len(raw_text) - len(stripped)]
            rest = stripped[len(marker.raw_marker) :]
            if rest.startswith(" "):
                rest = rest[1:]
            raw_text = leading_ws + rest
    text = _prepare_text(raw_text, normalisation=normalisation, apply_bidi=apply_bidi)
    style = "List Number" if marker is not None and marker.kind == "ordered" else "List Bullet"
    return document.add_paragraph(text, style=style)


def _add_failure(document: _DocxDocument, region: Region) -> _DocxParagraph:
    page_n = region.page_index + 1
    reason = region.failure_reason or "unknown"
    return document.add_paragraph(
        f"Transcription failed (page {page_n}): {reason}", style="Quote"
    )


def _add_figure(
    document: _DocxDocument,
    region: Region,
    *,
    pending_caption_text: dict[str, str],
) -> _DocxParagraph:
    page_n = region.page_index + 1
    text = f"Figure on page {page_n}"
    if region.group_id is not None and region.group_id in pending_caption_text:
        caption = pending_caption_text[region.group_id]
        if caption:
            text = f"{text}: {caption}"
    return document.add_paragraph(text, style="Normal")


def _add_table(
    document: _DocxDocument,
    region: Region,
    *,
    normalisation: NormalisationForm,
    rtl: bool,
) -> None:
    grid: TableGrid | None = region.table_grid
    if grid is None or grid.n_rows == 0 or grid.n_cols == 0:
        return
    if region.meta.get("v1_table_simplification") is True:
        marker = document.add_paragraph("(v1: merged cells flattened)", style="Normal")
        if rtl:
            _set_rtl_paragraph(marker)
    table = document.add_table(rows=grid.n_rows, cols=grid.n_cols)
    for r_idx, row in enumerate(grid.rows):
        for c_idx in range(grid.n_cols):
            cell_text = (
                normalise_text(row[c_idx].text, form=normalisation) if c_idx < len(row) else ""
            )
            table.cell(r_idx, c_idx).text = cell_text
    if rtl:
        _set_rtl_table(table)


__all__ = ["emit_docx"]
