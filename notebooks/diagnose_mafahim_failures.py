"""Build a Jupyter notebook that diagnoses why
``مفاهيم قرآنية في البناء والتنمية.pdf`` has 168/174 OCR failures and
why the .docx emitter crashes when writing the failure placeholders.

Run this script once:

    python notebooks/diagnose_mafahim_failures.py

It writes ``notebooks/diagnose_mafahim_failures.ipynb``. Open the
notebook and run cells top-to-bottom. The notebook is the artefact;
this Python file is just the builder (kept committed so the notebook
is reproducible).
"""

from __future__ import annotations

from pathlib import Path

import nbformat as nbf
from nbformat.v4 import new_code_cell, new_markdown_cell, new_notebook


PDF = (
    "/home/abdullah/Downloads/books-to-be-ocred/"
    "مفاهيم قرآنية في البناء والتنمية.pdf"
)


def cells() -> list:
    out: list = []

    out.append(
        new_markdown_cell(
            f"""# Diagnosing failures on `مفاهيم قرآنية في البناء والتنمية.pdf`

**Symptom (CLI)** — `summary: 174 pages, ok=6 failed=168`, then a
traceback in the docx emitter:

```
ValueError: All strings must be XML compatible: Unicode or ASCII,
no NULL bytes or control characters
```

This notebook reproduces both failures using the package as a library
and pinpoints the root causes.

PDF: `{PDF}`
"""
        )
    )

    # ---------------- 1. PDF inspection ---------------------------------
    out.append(new_markdown_cell("## 1. PDF metadata & page sampling"))
    out.append(
        new_code_cell(
            """\
import pypdfium2 as pdfium

PDF = r\"\"\"%s\"\"\"
doc = pdfium.PdfDocument(PDF)
print(f"pages: {len(doc)}")
for i in [0, 5, 10, 50, 100, 150, 173]:
    p = doc[i]
    tp = p.get_textpage()
    txt = tp.get_text_bounded()
    print(f"page {i:3d}: {p.get_width():.0f}x{p.get_height():.0f}  "
          f"native_chars={len(txt):4d}  preview={txt[:40]!r}")
"""
            % PDF
        )
    )
    out.append(
        new_markdown_cell(
            "**Reading**: most pages have **0 native chars** — scanned image "
            "pages. With `--force-ml` (and even without, since the validator "
            "would route empty pages to ML anyway) every page goes through "
            "the Surya OCR branch."
        )
    )

    # ---------------- 2. Render single page -----------------------------
    out.append(new_markdown_cell("## 2. Rasterise a few pages at 300 DPI"))
    out.append(
        new_code_cell(
            """\
from arabic_pdf_transcribe.pdf._pypdfium2_loader import open_pdf
from arabic_pdf_transcribe.layout._rasterise import rasterise_page

SAMPLE_PAGES = [0, 1, 5, 50, 100]

with open_pdf(PDF) as document:
    images = {}
    for idx in SAMPLE_PAGES:
        page = document[idx]
        try:
            img = rasterise_page(page, dpi=300)
        finally:
            page.close()
        images[idx] = img
        print(f"page {idx}: image {img.size}, mode {img.mode}")
"""
        )
    )

    # ---------------- 3. Run thread-mode pipeline on 5 pages ------------
    out.append(
        new_markdown_cell(
            "## 3. Thread-mode pipeline on pages 1–5\n\n"
            "`transcribe()` is the public library entry point. The CLI "
            "helpers `_build_layout`, `_build_ocr`, `_force_ml_validator` "
            "are reused so we hit the same code path the user invoked."
        )
    )
    out.append(
        new_code_cell(
            """\
from arabic_pdf_transcribe.pipeline import transcribe
from arabic_pdf_transcribe.cli import (
    _build_layout, _build_ocr, _force_ml_validator,
)

layout = _build_layout("doclayout-yolo", device="cuda")
ocr = _build_ocr("surya", device="cuda", disable_formula=True, batch_size=16)

events_t = []
def cb_t(page_index, total, event):
    events_t.append((page_index, event))

result_thread = transcribe(
    PDF,
    layout_detector=layout,
    ocr_transcriber=ocr,
    validator=_force_ml_validator,
    pages=range(0, 5),
    dpi=300,
    max_workers=1,
    progress=cb_t,
)

print(f"ok={result_thread.ok_pages} failed={result_thread.failed_pages}")
for region in result_thread.regions:
    print(f"page {region.page_index} role={region.role.value} "
          f"text_len={len(region.text)} fail={region.failure_reason!r}")
"""
        )
    )

    # ---------------- 4. Run ray-mode pipeline on 10 pages --------------
    out.append(
        new_markdown_cell(
            "## 4. Ray-mode pipeline on pages 1–10 (reproduce the OOM cascade)\n\n"
            "On a 6 GB GPU the layout actor + OCR actor each hold their own "
            "CUDA context. After the first OCR call PyTorch allocator state "
            "is fragmented and pages 6+ raise `torch.OutOfMemoryError`.\n\n"
            "**Note** — this cell loads two extra Ray actors and may take "
            "30–60 s the first time."
        )
    )
    out.append(
        new_code_cell(
            """\
import warnings
warnings.filterwarnings("ignore", category=FutureWarning)

# Free the thread-mode adapters first so the Ray actors don't have to
# share VRAM with them.
import gc, torch
del layout, ocr, result_thread
gc.collect()
torch.cuda.empty_cache()

from arabic_pdf_transcribe.cli import _build_ray_adapters

ray_layout, ray_ocr = _build_ray_adapters(
    device="cuda",
    layout_backend="doclayout-yolo",
    ocr_backend="surya",
    disable_formula=True,
    batch_size=16,
    num_gpus_per_actor=0.5,
)

events_r = []
def cb_r(page_index, total, event):
    events_r.append((page_index, event))

result_ray = transcribe(
    PDF,
    layout_detector=ray_layout,
    ocr_transcriber=ray_ocr,
    validator=_force_ml_validator,
    pages=range(0, 10),
    dpi=300,
    max_workers=1,
    progress=cb_r,
)

print(f"ok={result_ray.ok_pages} failed={result_ray.failed_pages}")
print()
for region in result_ray.regions:
    if region.failure_reason:
        first_line = region.failure_reason.splitlines()[0]
        print(f"page {region.page_index} FAIL: {first_line[:140]}...")
"""
        )
    )

    # ---------------- 5. Inspect the failure_reason string --------------
    out.append(
        new_markdown_cell(
            "## 5. Why python-docx rejects the failure reason\n\n"
            "lxml refuses control characters (`< 0x20` except `\\t \\n \\r`). "
            "`RayTaskError`'s `str()` includes ANSI colour escapes "
            "(`\\u001b[36m...\\u001b[39m`) which contain `\\x1b` (ESC) — a "
            "control character."
        )
    )
    out.append(
        new_code_cell(
            """\
import re

fail = next((r for r in result_ray.regions if r.failure_reason), None)
assert fail is not None, "Ray run did not produce a failure to inspect"
reason = fail.failure_reason

print("reason length:", len(reason))
print("first 200 chars:", reason[:200])
print()

ctrl = sorted({hex(ord(c)) for c in reason if ord(c) < 0x20 and c not in "\\t\\n\\r"})
print("control chars present:", ctrl)
print("ANSI ESC count:", reason.count("\\x1b"))

# Reproduce the python-docx crash on a single paragraph:
from docx import Document
doc = Document()
try:
    doc.add_paragraph(reason)
except ValueError as exc:
    print()
    print("python-docx rejection:", exc)
"""
        )
    )

    # ---------------- 6. Defensive sanitiser ----------------------------
    out.append(
        new_markdown_cell(
            "## 6. Proposed emitter fix — strip control chars before write\n\n"
            "`emit/docx.py:_add_failure` should sanitise the string before "
            "handing it to python-docx. The reason field is *diagnostic "
            "text*, never load-bearing, so dropping ANSI escapes is safe.\n\n"
            "Also worth doing: keep only the last informative line so the "
            "failure paragraph isn't a multi-KB stack trace."
        )
    )
    out.append(
        new_code_cell(
            """\
_CONTROL_RE = re.compile(r"[\\x00-\\x08\\x0b\\x0c\\x0e-\\x1f]")

def sanitise_for_docx(s: str, *, max_chars: int = 400) -> str:
    s = _CONTROL_RE.sub("", s)
    err_lines = [
        ln for ln in s.splitlines()
        if ln and ":" in ln and not ln.startswith(" ")
    ]
    pretty = err_lines[-1] if err_lines else s
    if len(pretty) > max_chars:
        pretty = pretty[: max_chars - 1] + "…"
    return pretty


cleaned = sanitise_for_docx(reason)
print("cleaned reason:", cleaned)

doc = Document()
doc.add_paragraph(
    f"Transcription failed (page {fail.page_index + 1}): {cleaned}",
    style="Quote",
)
doc.save("/tmp/sanitised_failure.docx")
print()
print("/tmp/sanitised_failure.docx written OK")
"""
        )
    )

    # ---------------- 7. CUDA OOM root cause ----------------------------
    out.append(
        new_markdown_cell(
            "## 7. Why Ray OOMs after page 5 on a 6 GB GPU\n\n"
            "* Thread mode: layout + OCR adapters live in **one** Python "
            "process and share a single CUDA context. Allocator "
            "fragmentation is bounded.\n"
            "* Ray mode (current design): layout actor and OCR actor are "
            "**separate** processes. Each opens its own CUDA context "
            "(~300–500 MB just to host the runtime), and each holds its own "
            "PyTorch caching allocator. Fixed overhead is doubled.\n"
            "* The trace shows PyTorch reports 4.48 GB allocated + 714 MB "
            "reserved-but-unallocated → fragmented allocator. Variable-"
            "sized activations (each page has a different number of "
            "detection slices) can't find a contiguous block.\n"
            "* `_handle_phase_exc` (`pipeline.py`) only flags the error as a "
            "CUDA-OOM when the exception class is exactly "
            "`torch.OutOfMemoryError`. A `RayTaskError(OCRTranscriptionError)` "
            "falls through to the generic `Exception` arm — so even the "
            "existing OOM-aware error message path is missed.\n\n"
            "**Fix avenues (least-to-most intrusive):**\n\n"
            "1. **Sanitise reasons in the emitter** — defensive, contract-"
            "level fix. Stops the `.docx` write from crashing regardless of "
            "what string the upstream produced.\n"
            "2. **Walk `__cause__`** in `pipeline._is_cuda_oom` so a wrapped "
            "torch OOM is recognised even through `RayTaskError` / "
            "`OCRTranscriptionError`. Then a `torch.cuda.empty_cache()` "
            "retry inside the actor becomes possible.\n"
            "3. **Single-actor Ray mode.** One actor that hosts both the "
            "layout and OCR adapters → one CUDA context, behaves like "
            "thread mode but with crash isolation. Also the right answer for "
            "≤6 GB GPUs.\n"
            "4. **Auto-retry on first OOM** — empty cache, halve "
            "`recognition_batch_size`, retry the same page once before "
            "marking as failed.\n"
        )
    )

    # ---------------- 8. Summary ----------------------------------------
    out.append(
        new_markdown_cell(
            "## 8. Summary\n\n"
            "| Symptom | Root cause | Suggested fix |\n"
            "|---|---|---|\n"
            "| 168/174 pages fail with `surya failed: CUDA out of memory` | "
            "Ray runs layout + OCR as **separate** CUDA processes on a 6 GB "
            "GPU; allocator fragmentation traps reserved memory after the "
            "first run | Single-actor co-resident mode, OR retry with "
            "`empty_cache()` + halved batch size when CUDA OOM is detected "
            "through the Ray wrapper |\n"
            "| `ValueError: All strings must be XML compatible…` from "
            "python-docx | `RayTaskError.__str__` embeds ANSI escape codes "
            "(`\\u001b[36m…`) which lxml rejects as control characters | "
            "Sanitise `failure_reason` in `emit/docx.py:_add_failure` "
            "(strip `\\x00-\\x1f` except `\\t \\n \\r`, optionally trim to "
            "the last exception line) |\n\n"
            "The first symptom is the upstream cause — fix it and the second "
            "goes away on its own — but the emitter sanitisation is still "
            "warranted as a contract: *no upstream string should be able to "
            "crash the writer*."
        )
    )

    return out


def main() -> None:
    nb = new_notebook(cells=cells())
    out = Path(__file__).with_suffix(".ipynb")
    out.write_text(nbf.writes(nb), encoding="utf-8")
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
